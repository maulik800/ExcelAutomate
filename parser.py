import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Inches,Pt
from datetime import date
from datetime import datetime
import sys
a = sys.argv[1]
row_no = int(a)-2
today = str(date.today())
dt= str(datetime.now())

df = pd.read_excel("F:\Atul\data1.xlsm")
last_row = df.iloc[row_no,:]
document = Document()
style = document.styles['Normal']
font = style.font
font.name = "Calibri"
font.size = Pt(14)
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# We set the paragraph formatting here
ref_para = document.add_paragraph('Ref.No.MI/2020-21/EMD/Refund')
ref_paragraph_format = ref_para.paragraph_format
ref_paragraph_format.space_after = Pt(0)

# We set the paragraph formatting here
dated_para = document.add_paragraph('Dated: '+today)
paragraph_format = dated_para.paragraph_format
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
dated_para.add_run("\n")



address_para = document.add_paragraph()
paragraph_format = dated_para.paragraph_format
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#The address para is added here
name_run = address_para.add_run(last_row[9]+",\n")
name_run.bold = True
add1_run  = address_para.add_run(last_row[10]+",\n")
add1_run.bold = True
add2_run  = address_para.add_run(last_row[11]+".\n")
add2_run.bold = True
add3_run  = address_para.add_run(last_row[12]+":\t"+str(last_row[13])+"\n")
add3_run.bold = True
add4_run  = address_para.add_run("E-mail Id:\t"+str(last_row[14])+"\n")
add4_run.bold = True


#The subject para is added here
subject_para = document.add_paragraph()
sub_head = subject_para.add_run("Subject:\t")
sub_head.bold = True
sub_con = subject_para.add_run("Refund EMD for "+last_row[6]+"\n")
sub_con.underline = True
ref = subject_para.add_run("Reference:\t")
ref.bold = True
ref_con = subject_para.add_run("Your NIT No "+str(last_row[4]))


body_para = document.add_paragraph()
body_run = body_para.add_run("Dear Sir,\n\n")
body1_run = body_para.add_run("With reference to above, we would like to inform you that we have submitted the above tender.\
 Our tender is not considered due as "+last_row[21]+".")


hence_para = document.add_paragraph()
hence_run = hence_para.add_run("Hence it is requested kindly arrange to refund our Earnest Money which is deposited along with the\
 tender documents in shape of "+str(last_row[15])+" wide no "+str(last_row[17])+" dated "+str(last_row[18])+" issued by "+last_row[16])


req_para = document.add_paragraph()
req_run = req_para.add_run("You are requested kindly arrange to refund our Earnest Money at your earliest possible or you may \
deposit the same to our account no. 11112223345 State Bank of India SMS Highway Branch Jaipur 302004 IFSC Code No. SBIN025316NM & oblige.\
\n\n\n\n\n")



thanks_para = document.add_paragraph()
thanks_run = req_para.add_run("Thanking You,\n\n\n")
faithfully_run = req_para.add_run("Your faithfully,\n")
metro_run = req_para.add_run("For Metro International.\n\n\n")
metro_run.bold = True

auth_run = req_para.add_run("(Authorised Signatory)")



document.save("F:/Atul/new/"+str(last_row[4])+".docx")
